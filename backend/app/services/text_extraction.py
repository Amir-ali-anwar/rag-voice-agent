import os
from pypdf import PdfReader
from docx import Document as DocxDocument


class TextExtractionService:

    SUPPORTED_FORMATS = {
        'text/plain': ['txt', 'md'],
        'application/pdf': ['pdf'],
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': ['docx'],
    }

    def extract_text(self, file_path: str, content_type: str):
        if not os.path.exists(file_path):
            raise FileNotFoundError(f'file not found :{file_path}')
        
        extension = self._get_extension(file_path)
        if extension == "text/plain" or extension in ['txt', 'md']:
            return self._extract_text_file(file_path)
        elif content_type == 'application/pdf' or extension == 'pdf':
            return self._extract_pdf_file(file_path)
        elif content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document' or extension == 'docx':
            return self._extract_word_file(file_path)
        else:
            raise ValueError(
                f"Unsupported file type: {content_type} or {extension}. "
                f"Supported formats: .txt, .md, .pdf, .docx"
            )

    def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text files"""

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                return text.strip()
            
        except UnicodeDecodeError:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
                return text.strip()

    def _extract_pdf_file(self, file_path: str) -> str:
        """Extract text from PDF files using pypdf"""

        try:
            reader = PdfReader(file_path)
            text_parts = []

            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text.strip())
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise ValueError(f"Error processing PDF file: {e}")

    
    def _extract_word_file(self, file_path: str) -> str:
        """Extract text from DOCX files using python-docx"""
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n\n".join(text_parts)
            
        except Exception as e:
            raise ValueError(f"Error processing DOCX file: {e}")
    
    def is_supported(self, content_type: str, file_path: str) -> bool:
        """Check if file type is supported"""
        if not content_type or not file_path:
            return False
        file_extension = self._get_extension(file_path)
        if content_type in self.SUPPORTED_FORMATS:
            return True

        for mime_type, extensions in self.SUPPORTED_FORMATS.items():
            if file_extension in extensions:
                return True    
   
    def _get_extension(self, file_path: str) -> str:
        """Get file extension without the dot"""
        return os.path.splitext(file_path.lower())[1][1:]
    
    
    